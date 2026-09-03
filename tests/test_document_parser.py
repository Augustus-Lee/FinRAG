"""文档解析器测试：Word / Excel / PDF（样本全部动态生成，不依赖外部文件）。"""

import pytest

from finrag.core.chunker import FinancialChunker
from finrag.core.document_parser import (
    ExcelParser,
    PyMuPDFParser,
    TextLikeParser,
    WordParser,
    create_parser,
)


# ---------------------------------------------------------------------------
# 工厂路由
# ---------------------------------------------------------------------------


def test_factory_routes_by_file_type():
    assert isinstance(create_parser("pdf"), PyMuPDFParser)
    assert isinstance(create_parser("PDF"), PyMuPDFParser)
    assert isinstance(create_parser("word"), WordParser)
    assert isinstance(create_parser("docx"), WordParser)
    assert isinstance(create_parser("excel"), ExcelParser)
    assert isinstance(create_parser("xlsx"), ExcelParser)
    assert isinstance(create_parser("md"), TextLikeParser)
    assert isinstance(create_parser(".md"), TextLikeParser)


def test_factory_rejects_unsupported_and_legacy():
    with pytest.raises(ValueError, match="不支持的文件类型"):
        create_parser("exe")
    with pytest.raises(ValueError, match="另存为"):
        create_parser("doc")
    with pytest.raises(ValueError, match="另存为"):
        create_parser("xls")


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------


def _make_docx(path, with_table: bool = True):
    from docx import Document

    doc = Document()
    doc.add_heading("产品费率", level=1)
    doc.add_paragraph("以下是费率说明。")
    doc.add_heading("收费明细", level=2)
    if with_table:
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "档位"
        table.cell(0, 1).text = "费率"
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "0.1%"
        table.cell(2, 0).text = "B"
        table.cell(2, 1).text = "0.05%"
    doc.save(str(path))
    return path


def test_word_parser_headings_and_full_table(tmp_path):
    path = _make_docx(tmp_path / "费率.docx")
    parsed = WordParser().parse(path)

    # 标题 → markdown 标题（切分器据此记录 section_path）
    assert "# 产品费率" in parsed.content
    assert "## 收费明细" in parsed.content
    # 表格完整数据行（回归：不允许只保留表头）
    assert "| 档位 | 费率 |" in parsed.content
    assert "| A | 0.1% |" in parsed.content
    assert "| B | 0.05% |" in parsed.content
    assert "以下是费率说明。" in parsed.content
    assert parsed.tables and parsed.tables[0]["headers"] == ["档位", "费率"]
    assert parsed.tables[0]["rows"] == 2


def test_word_parser_output_feeds_chunker(tmp_path):
    """Word 解析产物 → 切分器：标题层级 + 表格块语义完整传递。"""
    path = _make_docx(tmp_path / "费率.docx")
    parsed = WordParser().parse(path)
    chunks = FinancialChunker().split(parsed)

    table_chunks = [c for c in chunks if c.table_meta]
    assert len(table_chunks) == 1
    assert table_chunks[0].table_meta["headers"] == ["档位", "费率"]
    # 表格出现在"收费明细"标题之后 → section_path 应归属该标题
    assert table_chunks[0].section_path == "产品费率 / 收费明细"


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------


def _make_xlsx(path, n_rows: int = 2, extra_empty_sheet: bool = True):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "客户信息"
    ws.append(["姓名", "资产", "风险等级"])
    for i in range(n_rows):
        ws.append([f"客户{i}", 1000000 + i, "低"])
    if extra_empty_sheet:
        wb.create_sheet("空表")
    wb.save(str(path))
    return path


def test_excel_parser_sheet_as_heading_and_table(tmp_path):
    path = _make_xlsx(tmp_path / "客户.xlsx")
    parsed = ExcelParser().parse(path)

    assert "## 客户信息" in parsed.content  # sheet 名 → 标题
    assert "| 姓名 | 资产 | 风险等级 |" in parsed.content
    assert "| 客户0 | 1000000 | 低 |" in parsed.content
    assert "| 客户1 | 1000001 | 低 |" in parsed.content
    # 空 sheet 跳过
    assert "空表" not in parsed.content
    assert parsed.tables[0]["headers"] == ["姓名", "资产", "风险等级"]
    assert parsed.tables[0]["rows"] == 2


def test_excel_parser_splits_large_sheet_into_blocks(tmp_path):
    """大表按行分块：每块重复表头，避免单 chunk 撑爆 embedding 上下文。"""
    path = _make_xlsx(tmp_path / "大表.xlsx", n_rows=120, extra_empty_sheet=False)
    parsed = ExcelParser(max_rows=50).parse(path)

    # 120 数据行 / 50 行每块 → 3 块，各含表头 + 分隔行
    assert len(parsed.tables) == 3
    assert [t["rows"] for t in parsed.tables] == [50, 50, 20]
    # 每块表头重复，保证块自包含
    assert parsed.content.count("| 姓名 | 资产 | 风险等级 |") == 3
    assert "| 客户119 | 1000119 | 低 |" in parsed.content  # 末行不丢

    # 切分器视角：3 个独立表格 chunk，同一 section（sheet 名）
    chunks = FinancialChunker().split(parsed)
    table_chunks = [c for c in chunks if c.table_meta]
    assert len(table_chunks) == 3
    assert {c.section_path for c in table_chunks} == {"客户信息"}


def test_excel_parser_escapes_pipe_in_cells(tmp_path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "特殊"
    ws.append(["说明"])
    ws.append(["A|B"])
    path = tmp_path / "特殊.xlsx"
    wb.save(str(path))

    parsed = ExcelParser().parse(path)
    assert "A\\|B" in parsed.content  # | 转义，不破坏 markdown 表结构


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def _make_pdf_with_table(path):
    """fitz 生成含网格表格的 PDF：画线成表 + 单元格文字（ASCII 避免字体依赖）。"""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 50), "Rate Table Introduction")

    # 3 行 x 2 列网格：竖线 x=72,148,224；横线 y=80,120,160,200（含底边）
    for x in (72, 148, 224):
        page.draw_line(fitz.Point(x, 80), fitz.Point(x, 200))
    for y in (80, 120, 160, 200):
        page.draw_line(fitz.Point(72, y), fitz.Point(224, y))
    cells = [("Tier", "Rate"), ("A", "0.1%"), ("B", "0.05%")]
    for r, (c0, c1) in enumerate(cells):
        page.insert_text((80, 80 + 40 * r + 25), c0)
        page.insert_text((156, 80 + 40 * r + 25), c1)

    page.insert_text((72, 230), "End of document")  # 表格（y≤200）下方
    doc.save(str(path))
    doc.close()
    return path


def test_pdf_parser_extracts_text_and_full_table(tmp_path):
    path = _make_pdf_with_table(tmp_path / "rate.pdf")
    parsed = PyMuPDFParser().parse(path)

    # 正文文本（表格外的普通文本块）
    assert "Rate Table Introduction" in parsed.content
    assert "End of document" in parsed.content
    # 表格完整数据行（回归：旧实现只有表头、无数据行）
    assert "| Tier | Rate |" in parsed.content
    assert "| A | 0.1% |" in parsed.content
    assert "| B | 0.05% |" in parsed.content
    assert parsed.tables and parsed.tables[0]["rows"] == 2
    # 表格文字不重复出现（与表格 bbox 重叠的文本块已过滤）
    assert parsed.content.count("0.05%") == 1


def test_pdf_parser_table_between_sections(tmp_path):
    """表格前后都有文本 → 交错输出保持文档顺序（表格在原文位置）。"""
    path = _make_pdf_with_table(tmp_path / "rate.pdf")
    parsed = PyMuPDFParser().parse(path)

    intro = parsed.content.index("Rate Table Introduction")
    table = parsed.content.index("| Tier | Rate |")
    end = parsed.content.index("End of document")
    assert intro < table < end
