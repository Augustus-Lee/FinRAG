"""摄入流水线测试：解析器按 doc.file_type 动态选择（回归 bug：曾硬编码 md 解析器）。"""

import pytest

from finrag.core.bm25 import BM25Index
from finrag.core.chunker import FinancialChunker
from finrag.db.session import SessionLocal
from finrag.models import KBCategory, KBChunk, KBDocument
from finrag.pipelines.ingest import IngestPipeline


class _FakeEmbedding:
    """固定 8 维向量，避免加载真实模型。"""

    def embed(self, texts: list[str]) -> list[list[float]]:
        return [[0.1] * 8 for _ in texts]

    def embed_query(self, text: str) -> list[float]:
        return [0.1] * 8


class _FakeVectorStore:
    def __init__(self) -> None:
        self.points: list = []

    def upsert(self, points) -> None:
        self.points.extend(points)

    def delete_by_ids(self, ids) -> None:
        self.points = [p for p in self.points if p.id not in ids]


def _make_pipeline() -> tuple[IngestPipeline, _FakeVectorStore]:
    vs = _FakeVectorStore()
    pipeline = IngestPipeline(
        chunker=FinancialChunker(),
        embedding=_FakeEmbedding(),
        vector_store=vs,
        bm25_index=BM25Index(),
    )
    return pipeline, vs


def _make_docx(path):
    from docx import Document

    doc = Document()
    doc.add_heading("产品费率", level=1)
    doc.add_paragraph("费率说明正文。")
    table = doc.add_table(rows=3, cols=2)
    for r, row in enumerate([("档位", "费率"), ("A", "0.1%"), ("B", "0.05%")]):
        table.cell(r, 0).text = row[0]
        table.cell(r, 1).text = row[1]
    doc.save(str(path))
    return path


def _cleanup(db, doc_id: int, cat_id: int) -> None:
    db.query(KBChunk).filter(KBChunk.doc_id == doc_id).delete(synchronize_session=False)
    db.query(KBDocument).filter(KBDocument.id == doc_id).delete(synchronize_session=False)
    db.query(KBCategory).filter(KBCategory.id == cat_id).delete(synchronize_session=False)
    db.commit()


def test_ingest_selects_parser_by_doc_file_type(tmp_path):
    """file_type=docx → 必须用 WordParser（旧 bug：硬编码 md 解析器把 docx 当文本读成乱码）。"""
    db = SessionLocal()
    try:
        cat = KBCategory(name="解析测试库", owner_id=1)
        db.add(cat)
        db.commit()
        db.refresh(cat)
        path = _make_docx(tmp_path / "费率.docx")
        doc = KBDocument(
            kb_id=cat.id, name="费率.docx", file_type="docx", file_path=str(path), owner_id=1
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)

        pipeline, vs = _make_pipeline()
        result = pipeline.ingest_document(db, doc, str(path))

        assert result.status == "ready"
        assert result.chunk_count > 0
        assert doc.status == "ready"

        chunks = db.query(KBChunk).filter(KBChunk.doc_id == doc.id).all()
        assert len(chunks) == result.chunk_count
        # 解析器正确性的直接证据：section_path 来自 Word 标题、表格 chunk 有 table_meta
        assert any(c.section_path == "产品费率" for c in chunks)
        table_chunks = [c for c in chunks if c.table_meta]
        assert len(table_chunks) == 1
        assert table_chunks[0].table_meta["headers"] == ["档位", "费率"]
        # 确定性 vector_id：Qdrant point / BM25 doc / KBChunk.vector_id 三方一致
        assert all(c.vector_id for c in chunks)
        assert [p.id for p in vs.points] == [c.vector_id for c in chunks]
        # BM25 索引同步可查
        assert pipeline._bm25.search("费率", top_k=3)

        _cleanup(db, doc.id, cat.id)
    finally:
        db.close()


def test_ingest_failed_on_unsupported_file_type(tmp_path):
    db = SessionLocal()
    try:
        cat = KBCategory(name="解析测试库2", owner_id=1)
        db.add(cat)
        db.commit()
        db.refresh(cat)
        path = tmp_path / "fake.exe"
        path.write_bytes(b"\x00\x01")
        doc = KBDocument(
            kb_id=cat.id, name="fake.exe", file_type="exe", file_path=str(path), owner_id=1
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)

        pipeline, _ = _make_pipeline()
        with pytest.raises(ValueError, match="不支持的文件类型"):
            pipeline.ingest_document(db, doc, str(path))

        # 失败状态落库，错误信息可追溯
        db.refresh(doc)
        assert doc.status == "failed"
        assert "不支持的文件类型" in doc.error_msg

        _cleanup(db, doc.id, cat.id)
    finally:
        db.close()
