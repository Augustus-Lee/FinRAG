"""文档解析器：Markdown / 纯文本 / PDF / Word / Excel。

策略：不重造解析引擎（避免 RAGFlow DeepDoc 的复杂度），组合成熟库——
- md / txt：直接读取文本
- PDF：PyMuPDF 文本 + 表格结构化抽取（完整数据行，按原文位置交错）
- Word(.docx)：python-docx，标题→markdown 标题、表格→markdown 表
- Excel(.xlsx)：openpyxl，sheet→标题 + markdown 表（大表按行分块）

统一输出 markdown 格式 content：切分器（FinancialChunker）原生支持标题层级
（section_path 溯源）与表格块（table_meta + 整体成 chunk），解析器无需感知切分细节。

扫描件 OCR、PDF 矢量图形（柱状图/折线图）内容抽取为 Out of scope：
图形无文本可抽，其标题/图注作为普通文本自然流入 content。
"""

import re
from abc import ABC, abstractmethod
from datetime import date, datetime
from pathlib import Path

from finrag.core.chunker import ParsedDocument
from finrag.logging import get_logger

logger = get_logger("finrag.document_parser")

_WORD_TYPES = {"word", "docx"}
_EXCEL_TYPES = {"excel", "xlsx"}
_PDF_TYPES = {"pdf", "pdfs"}
_LEGACY_HINT = "旧格式请先另存为 .docx / .xlsx 再上传"

# 支持类型：word/docx/excel/xlsx 为别名对，file_type 统一按扩展名规范化存储
_SUPPORTED = {"md", "txt", "markdown", "pdf", "pdfs"} | _WORD_TYPES | _EXCEL_TYPES

# Word 标题样式：兼容英文 "Heading 1" 与中文 Word "标题 1"
_HEADING_STYLE_RE = re.compile(r"^(?:Heading|标题)\s*([1-9])")


def _cell_text(value: object) -> str:
    """单元格值规整为安全文本：None→空串，日期 ISO 化，| 与换行转义避免破坏 markdown 表。"""
    if value is None:
        return ""
    if isinstance(value, (datetime, date)):
        text = value.isoformat(sep=" ", timespec="seconds") if isinstance(value, datetime) else value.isoformat()
    else:
        text = str(value)
    return text.replace("|", "\\|").replace("\r\n", " ").replace("\n", " ").strip()


def _md_table(rows: list[list[str]]) -> str:
    """二维数据 → 完整 markdown 表格（表头 + 分隔行 + 全部数据行）。"""
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    norm = [list(r) + [""] * (width - len(r)) for r in rows]
    lines = [
        "| " + " | ".join(norm[0]) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    lines.extend("| " + " | ".join(r) + " |" for r in norm[1:])
    return "\n".join(lines)


def _table_meta(rows: list[list[str]]) -> dict:
    return {"headers": rows[0] if rows else [], "rows": max(0, len(rows) - 1)}


class DocumentParser(ABC):
    @abstractmethod
    def parse(self, path: str | Path) -> ParsedDocument: ...


class TextLikeParser(DocumentParser):
    """md / txt 解析：直接读取为 markdown 文本。"""

    def parse(self, path: str | Path) -> ParsedDocument:
        p = Path(path)
        content = p.read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(title=p.stem, content=content)


class WordParser(DocumentParser):
    """Word(.docx) 解析：python-docx，延迟导入。

    按 body 元素顺序迭代（w:p 段落 / w:tbl 表格），保持原文顺序；
    标题样式（Heading N / 标题 N）→ markdown # 标题，表格 → 完整 markdown 表。
    """

    def parse(self, path: str | Path) -> ParsedDocument:
        from docx import Document as DocxDocument  # python-docx，延迟导入
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        p = Path(path)
        doc = DocxDocument(str(p))
        parts: list[str] = []
        tables: list[dict] = []

        # body.iterchildren() 保持段落与表格的文档顺序（doc.paragraphs 会丢失表格位置）
        for child in doc.element.body.iterchildren():
            if child.tag.endswith("}p"):
                para = Paragraph(child, doc)
                text = para.text.strip()
                if not text:
                    continue
                m = _HEADING_STYLE_RE.match(para.style.name or "")
                if m:
                    level = int(m.group(1))
                    parts.append(f"{'#' * min(level, 6)} {text}")
                else:
                    parts.append(text)
            elif child.tag.endswith("}tbl"):
                table = Table(child, doc)
                rows = [[_cell_text(c.text) for c in row.cells] for row in table.rows]
                md = _md_table(rows)
                if md:
                    parts.append(md)
                    tables.append(_table_meta(rows))

        content = "\n\n".join(parts)
        logger.info("word_parsed", path=str(p), parts=len(parts), tables=len(tables))
        return ParsedDocument(title=p.stem, content=content, tables=tables)


class ExcelParser(DocumentParser):
    """Excel(.xlsx) 解析：openpyxl，延迟导入。

    每个 sheet → `## {sheet名}` 标题 + markdown 表（首个非空行作表头）；
    数据行超过 max_rows 时按行分块（每块重复表头），避免大表撑爆 chunk/embedding。
    """

    def __init__(self, max_rows: int = 50) -> None:
        if max_rows < 1:
            raise ValueError("max_rows 必须 ≥ 1")
        self._max_rows = max_rows

    def parse(self, path: str | Path) -> ParsedDocument:
        from openpyxl import load_workbook  # 延迟导入

        p = Path(path)
        # read_only 流式读取（大文件不整体载入）；data_only 取公式缓存值（非公式本身）
        wb = load_workbook(str(p), read_only=True, data_only=True)
        parts: list[str] = []
        tables: list[dict] = []
        try:
            for sheet in wb.worksheets:
                rows = [
                    [_cell_text(c) for c in row]
                    for row in sheet.iter_rows(values_only=True)
                ]
                rows = [r for r in rows if any(cell.strip() for cell in r)]
                if not rows:
                    continue
                parts.append(f"## {sheet.title}")
                # 大表按行分块：每块 ≤ max_rows 数据行，重复表头保证每块自包含
                data = rows[1:]
                for start in range(0, len(data), self._max_rows):
                    block = [rows[0]] + data[start : start + self._max_rows]
                    parts.append(_md_table(block))
                    tables.append(_table_meta(block))
        finally:
            wb.close()

        content = "\n\n".join(parts)
        logger.info("excel_parsed", path=str(p), sheets=len(wb.sheetnames), tables=len(tables))
        return ParsedDocument(title=p.stem, content=content, tables=tables)


class PyMuPDFParser(DocumentParser):
    """PDF 解析：PyMuPDF，延迟导入。文本 + 表格结构化抽取（完整数据行，按原文位置交错）。"""

    def parse(self, path: str | Path) -> ParsedDocument:
        import fitz  # PyMuPDF，延迟导入

        p = Path(path)
        doc = fitz.open(p)
        parts: list[str] = []
        tables: list[dict] = []
        for page in doc:
            # 表格 bbox：用于过滤与表格重叠的文本块（表格文字已由 extract() 结构化，
            # 不再重复出现在正文里）
            found = page.find_tables()
            table_items = [
                (tbl.bbox, [[_cell_text(c) for c in row] for row in (tbl.extract() or [])])
                for tbl in found.tables
            ]

            def _in_table(bbox: tuple) -> bool:
                x0, y0, x1, y1 = bbox
                return any(
                    not (x1 < tb[0] or x0 > tb[2] or y1 < tb[1] or y0 > tb[3])
                    for tb, _ in table_items
                )

            # 文本块（type 0）与表格按 y 坐标排序交错 → 表格出现在原文位置
            items: list[tuple[float, float, str]] = []
            for b in page.get_text("blocks"):
                x0, y0, x1, y1, text, _, block_type = b[:7]
                if block_type == 0 and text.strip() and not _in_table((x0, y0, x1, y1)):
                    items.append((y0, x0, text.strip()))
            for bbox, rows in table_items:
                if any(any(c.strip() for c in r) for r in rows):
                    items.append((bbox[1], bbox[0], _md_table(rows)))
                    tables.append(_table_meta(rows))

            items.sort(key=lambda t: (t[0], t[1]))
            for _, _, text in items:
                parts.append(text)

        content = "\n\n".join(parts)
        logger.info("pdf_parsed", path=str(p), pages=len(doc), tables=len(tables))
        return ParsedDocument(title=p.stem, content=content, tables=tables)


def create_parser(file_type: str) -> DocumentParser:
    """按文件类型创建解析器。"""
    ft = file_type.lower().lstrip(".")
    if ft in {"doc", "xls"}:
        raise ValueError(f"不支持旧格式 .{ft}：{_LEGACY_HINT}")
    if ft not in _SUPPORTED:
        raise ValueError(f"不支持的文件类型: {file_type}（支持 {sorted(_SUPPORTED - {'pdfs'})}）")
    if ft in _PDF_TYPES:
        return PyMuPDFParser()
    if ft in _WORD_TYPES:
        return WordParser()
    if ft in _EXCEL_TYPES:
        return ExcelParser()
    return TextLikeParser()
